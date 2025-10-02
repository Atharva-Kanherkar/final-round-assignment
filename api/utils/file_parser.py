"""Utility for parsing resume files (PDF, DOCX, TXT)."""
import io
import logging
from typing import Tuple
from fastapi import UploadFile, HTTPException

logger = logging.getLogger(__name__)

# Lazy imports to prevent startup failure if dependencies not installed
def _get_pdf_reader():
    try:
        from pypdf import PdfReader
        return PdfReader
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="PDF parsing not available. Install pypdf package."
        )

def _get_docx_document():
    try:
        from docx import Document
        return Document
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="DOCX parsing not available. Install python-docx package."
        )


class FileParser:
    """Parse resume and job description files."""

    ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}
    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB

    @staticmethod
    async def parse_file(file: UploadFile) -> str:
        """
        Parse uploaded file and extract text content.

        Args:
            file: Uploaded file from FastAPI

        Returns:
            Extracted text content

        Raises:
            HTTPException: If file format is not supported or parsing fails
        """
        try:
            # Check file extension
            filename = file.filename.lower() if file.filename else ""
            ext = None
            for allowed_ext in FileParser.ALLOWED_EXTENSIONS:
                if filename.endswith(allowed_ext):
                    ext = allowed_ext
                    break

            if not ext:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file format. Allowed: {', '.join(FileParser.ALLOWED_EXTENSIONS)}"
                )

            # Read file content
            content = await file.read()

            # Check file size
            if len(content) > FileParser.MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=400,
                    detail=f"File too large. Maximum size: {FileParser.MAX_FILE_SIZE / 1024 / 1024}MB"
                )

            # Parse based on extension
            if ext == '.pdf':
                text = FileParser._parse_pdf(content)
            elif ext in ['.docx', '.doc']:
                text = FileParser._parse_docx(content)
            elif ext == '.txt':
                text = content.decode('utf-8', errors='ignore')
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file format: {ext}"
                )

            # Validate extracted text
            text = text.strip()
            if len(text) < 50:
                raise HTTPException(
                    status_code=400,
                    detail="Extracted text is too short (minimum 50 characters). Please check your file."
                )

            logger.info(f"Successfully parsed {ext} file: {filename} ({len(text)} characters)")
            return text

        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error parsing file {file.filename}: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to parse file: {str(e)}"
            )

    @staticmethod
    def _parse_pdf(content: bytes) -> str:
        """
        Parse PDF file content.

        Args:
            content: PDF file bytes

        Returns:
            Extracted text
        """
        try:
            PdfReader = _get_pdf_reader()
            pdf_file = io.BytesIO(content)
            pdf_reader = PdfReader(pdf_file)

            text_parts = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            return '\n'.join(text_parts)
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    def _parse_docx(content: bytes) -> str:
        """
        Parse DOCX file content.

        Args:
            content: DOCX file bytes

        Returns:
            Extracted text
        """
        try:
            Document = _get_docx_document()
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            return '\n'.join(text_parts)
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
